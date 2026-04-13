import pandas as pd
from docx import Document
import os

EXCEL_PATH = "data/ProbioCheck-50-Cepas-Semilla.xlsx"
WORD_PATH = "data/ProbioCheck-CBD.docx"

def extraer_excel(ruta_archivo=EXCEL_PATH):
    """
    Lee todas las hojas del Excel y las carga en memoria.
    Retorna un diccionario donde las claves son los nombres de las hojas
    y los valores son DataFrames de Pandas.
    """
    if not os.path.exists(ruta_archivo):
        raise FileNotFoundError(f"Error: No se encontro el archivo Excel en {ruta_archivo}")
    
    print(f"Extrayendo datos del Excel: {ruta_archivo}...")
    
    # sheet_name=None lee todas las pestanas de una vez
    diccionario_hojas = pd.read_excel(ruta_archivo, sheet_name=None, engine='openpyxl', header=1)
    
    print(f"Hojas extraidas: {list(diccionario_hojas.keys())}")
    return diccionario_hojas

def extraer_word(ruta_archivo=WORD_PATH):
    """
    Lee el documento de Word y extrae tanto el texto libre como las tablas.
    Retorna un diccionario con listas de párrafos y datos tabulares.
    """
    if not os.path.exists(ruta_archivo):
        raise FileNotFoundError(f"Error: No se encontro el archivo Word en {ruta_archivo}")
    
    print(f"Extrayendo datos del Word: {ruta_archivo}...")
    doc = Document(ruta_archivo)
    
    # ignoramos lineas en blanco
    parrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # las tablas del Word tienen los esquemas y scores
    datos_tablas = []
    for tabla in doc.tables:
        matriz_tabla = []
        for fila in tabla.rows:
            textos_celdas = [celda.text.strip() for celda in fila.cells]
            matriz_tabla.append(textos_celdas)
        datos_tablas.append(matriz_tabla)
        
    print(f"Extraidos {len(parrafos)} parrafos y {len(datos_tablas)} tablas.")
    
    return {
        "parrafos": parrafos,
        "tablas": datos_tablas
    }

if __name__ == "__main__":
    print("--- Iniciando prueba de extraccion ---")
    datos_excel = extraer_excel()
    datos_word = extraer_word()
    
    # vista rapida de las columnas de la primera hoja
    hoja_principal = list(datos_excel.keys())[0]
    print(f"\nColumnas de '{hoja_principal}':")
    print(datos_excel[hoja_principal].columns.tolist())
    print("--------------------------------------")